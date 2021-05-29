from dataclasses import dataclass
from typing import *
import time
import docx
import re


@dataclass
class Reference(object):
    number_in_text: int
    actual_number: int


@dataclass
class Link(object):
    text: str
    number_in_text: int
    actual_number: int
    p_number: int
    p2change: int
    style: docx.Document


class EditParagraph(object):

    def __init__(self, paragraph, p_number, refs):
        self._text = paragraph.text
        self._style = paragraph.style
        self._p_number = p_number
        self.refs = refs
        self.__edited = False

    @property
    def text(self):
        return self._text

    @text.setter
    def text(self, text: str):
        self._text = text

    @property
    def style(self):
        return self._style

    @style.setter
    def style(self, val: Any):
        raise AttributeError('Impossible to set a style')

    @property
    def p_number(self):
        return self._p_number

    @p_number.setter
    def p_number(self, val: Any):
        raise AttributeError('Impossible to set number of a paragraph')

    @property
    def refs(self):
        return self._refs

    @refs.setter
    def refs(self, refs: List[Tuple[int, int]]):
        self._refs = [
            Reference(num_in_text, actual_num)
            for num_in_text, actual_num in refs
        ]

    def edit_text(self):

        if self.__edited:
            return self.text

        def strlist2numlist(array: List[str]) -> List[int]:
            return [
                int(elem)
                for elem in array
            ]

        def replace_ref(num: int) -> int:
            for ref in paragraph.refs:
                if ref.number_in_text == num:
                    return ref.actual_number

        replaced_mark = 'THIS-TEXT-HAS-BEEN-REPLACED-WITH-REFFORMER-HAVE-A-NICE-DAY:*'
        text = self.text
        refs = re.findall(r'\[\d*,?\s?\d*,?\s?\d*,?\s?\d*,?\s?\d*,?\s?\d*,?\s?\d*,?\s?\d*]', text)

        numeric_refs = [strlist2numlist(
            re.findall(r'\d\d*', string)
        ) for string in refs]

        fixed_numeric_refs = [
            [
                replace_ref(num)
                for num in ref
            ]
            for ref in numeric_refs
        ]

        fixed_refs = [
            str(fixed_numeric_ref)
            for fixed_numeric_ref in fixed_numeric_refs
        ]

        for ref, fref in zip(refs, fixed_refs):
            text = text.replace(ref, fref[:1] + replaced_mark + fref[1:])

        # Huck-fuck code :)
        text = text.replace(replaced_mark, '')

        self.text = text
        self.__edited = True

        return self.text


if __name__ == '__main__':
    path = '/home/user/path/to/my/thesis.docx'
    path = '/home/user/Documents/Thesis/BTS-41-zachot-osen-2020-Zabolotniy-LitObzor_1_test.docx'
    print('RefFormer is starting...')

    if path == '/home/user/path/to/my/thesis.docx':
        path = input('Enter path to the document: ')
    start = time.time()
    doc = docx.Document(path)
    print('Document found. Reading...')
    ref_section: bool = False
    paragraphs = list()
    refs_count = 0
    all_numbers = dict()
    all_links = list()

    for paragraph, p_number in zip(doc.paragraphs, range(len(doc.paragraphs))):
        refs = re.findall(r'\[\d*,?\s?\d*,?\s?\d*,?\s?\d*,?\s?\d*,?\s?\d*,?\s?\d*,?\s?\d*]', paragraph.text)
        if refs:
            nums = list()

            for ref in refs:
                nums += re.findall(r'\d+', ref)
                nums = [int(num) for num in nums]

            for num in nums:

                if num in all_numbers:
                    continue
                else:
                    all_numbers.update({
                        num: len(all_numbers.values()) + 1
                    })

            paragraphs.append(
                EditParagraph(
                    paragraph,
                    p_number,
                    [
                        (num, all_numbers[num])
                        for num in nums
                    ]
                )
            )

        if 'References section' in paragraph.text:
            ref_section = True

        if ref_section:
            links = re.findall(r'^\s*\d\d*\.\s', paragraph.text)

            if links:
                lnum = int(
                    re.findall(r'\d\d*', links[0])[0]
                )

                for i in range(len(paragraph.text)):
                    if paragraph.text[i] == '.':
                        text = paragraph.text[i + 2:]
                        break
                if not text:
                    raise ValueError(f'Check the paragraph:\n{paragraph.text}')

                all_links.append(
                    Link(
                        f'{all_numbers[lnum]}. {text}',
                        lnum,
                        all_numbers[lnum],
                        p_number,
                        p_number + all_numbers[lnum] - lnum,
                        paragraph.style
                    )
                )

    for paragraph in paragraphs:
        paragraph.edit_text()

    for paragraph in paragraphs:
        doc.paragraphs[paragraph.p_number].text = paragraph.text
        doc.paragraphs[paragraph.p_number].style = paragraph.style

    for link in all_links:
        doc.paragraphs[link.p2change].text = link.text
        doc.paragraphs[link.p2change].style = link.style

    doc.save(path)

    print(f'RefFormer has finished working. Check carefully if everything is correct.\nRUNTIME: {time.time() - start}')
